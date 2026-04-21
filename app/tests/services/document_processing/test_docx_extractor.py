import zipfile
from pathlib import Path

from docx import Document

from app.services.document_processing.docx_extractor import DocxExtractor


def _inject_rendered_page_breaks(doc_path: Path, count: int = 1):
    with zipfile.ZipFile(doc_path, "a") as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        replacement = "".join("<w:lastRenderedPageBreak/>" for _ in range(count)) + "</w:r></w:p>"
        document_xml = document_xml.replace("</w:r></w:p>", replacement, 1)
        archive.writestr("word/document.xml", document_xml)


def test_extract_docx_paragraphs_in_order(tmp_path: Path):
    doc_path = tmp_path / "paragraphs.docx"
    document = Document()
    document.add_heading("Employee Handbook", level=1)
    document.add_paragraph("Welcome to the company.")
    document.add_paragraph("Please read the policies carefully.")
    document.save(doc_path)
    _inject_rendered_page_breaks(doc_path)

    result = DocxExtractor().extract_docx(str(doc_path))

    assert result["metadata"]["extraction_method"] == "docx_native_paginated"
    assert result["metadata"]["ocr_engine"] == "none"
    assert result["metadata"]["page_numbering"] == "physical_docx"
    assert result["metadata"]["total_pages"] == 2
    assert result["pages"][0]["text_content"].startswith("Employee Handbook")
    assert "Welcome to the company." in result["pages"][0]["text_content"]
    assert result["pages"][0]["page_number"] == 1
    assert result["document_structure"]["title"] == "Employee Handbook"


def test_extract_docx_tables_are_structured(tmp_path: Path):
    doc_path = tmp_path / "tables.docx"
    document = Document()
    document.add_paragraph("Benefits")
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Policy"
    table.cell(0, 1).text = "Owner"
    table.cell(1, 0).text = "Leave"
    table.cell(1, 1).text = "HR"
    table.cell(2, 0).text = "Security"
    table.cell(2, 1).text = "IT"
    document.save(doc_path)
    _inject_rendered_page_breaks(doc_path)

    result = DocxExtractor().extract_docx(str(doc_path))

    assert len(result["tables"]) == 1
    assert result["tables"][0]["headers"] == ["Policy", "Owner"]
    assert result["tables"][0]["data"][0]["Policy"] == "Leave"
    assert result["tables"][0]["page"] == 1
    assert "Policy | Owner" in result["pages"][0]["text_content"]
    assert "Security | IT" in result["pages"][0]["text_content"]


def test_extract_empty_docx_returns_compatible_structure(tmp_path: Path):
    doc_path = tmp_path / "empty.docx"
    Document().save(doc_path)

    result = DocxExtractor().extract_docx(str(doc_path))

    assert result["pages"] == []
    assert result["tables"] == []
    assert result["document_structure"]["title"] == ""
    assert result["metadata"]["total_pages"] == 0
    assert result["metadata"]["page_numbering"] == "physical_docx"


def test_extract_docx_without_rendered_breaks_fails_clearly(tmp_path: Path):
    doc_path = tmp_path / "no-breaks.docx"
    document = Document()
    document.add_heading("Employee Handbook", level=1)
    document.add_paragraph("Welcome to the company.")
    document.save(doc_path)

    try:
        DocxExtractor().extract_docx(str(doc_path))
        assert False, "Expected DOCX extraction to fail without rendered page markers"
    except ValueError as exc:
        assert "no rendered page markers" in str(exc)
